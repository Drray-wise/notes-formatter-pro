#!/usr/bin/env python3
"""
Notes Formatter Pro — Cloud / Web Server Edition
Developed by Mr Wise

Deploy free on Render.com:
  1. Push this folder to a GitHub repo
  2. New Web Service on render.com → connect repo → Done

Or run locally:
  pip install flask groq python-docx gunicorn
  python cloud_server.py
"""

import os, re, json, io
from flask import Flask, request, jsonify, Response, send_file, session
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__)
# Secret key for session cookies — change this to anything random if you want
app.secret_key = os.environ.get("SECRET_KEY", "nfpro-mr-wise-2025-secret")

DEFAULT_SETTINGS = {
    "theme": "midnight", "font_family": "Inter",
    "font_size": 13, "style": "Clean"
}

# ── Word export helpers ───────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#")); tcPr.append(shd)

def add_banner(doc, text):
    tbl = doc.add_table(rows=1, cols=1); tbl.style = "Table Grid"
    cell = tbl.cell(0, 0); set_cell_bg(cell, "6C47FF"); cell.width = Inches(6.5)
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("  " + text); run.bold = True
    run.font.color.rgb = RGBColor(255,255,255)
    run.font.size = Pt(14); run.font.name = "Arial"
    doc.add_paragraph("")

def add_docx_table(doc, headers, rows, font_name, font_size):
    if not rows: return
    ncols = max(len(headers) if headers else 1, max((len(r) for r in rows), default=1))
    tbl = doc.add_table(rows=1+len(rows), cols=ncols); tbl.style = "Table Grid"
    for ci, hdr in enumerate(headers):
        cell = tbl.cell(0, ci); set_cell_bg(cell, "6C47FF")
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hdr); run.bold = True
        run.font.color.rgb = RGBColor(255,255,255)
        run.font.size = Pt(10); run.font.name = font_name
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row[:ncols]):
            p = tbl.cell(ri+1,ci).paragraphs[0]
            run = p.add_run(val); run.font.size = Pt(10); run.font.name = font_name
    doc.add_paragraph("")

def clean_markdown(text):
    text = re.sub(r"^#{1,6}\s*","", text, flags=re.MULTILINE)
    for pat, repl in [
        (r"\*\*(.+?)\*\*",r"\1"),(r"\*(.+?)\*",r"\1"),
        (r"__(.+?)__",r"\1"),(r"_(.+?)_",r"\1"),
        (r"`{1,3}([^`]*)`{1,3}",r"\1"),(r"\[(.+?)\]\(.+?\)",r"\1"),
        (r"^\s*>\s*","",),(r"^[-*]{3,}\s*$",""),
    ]:
        text = re.sub(pat, repl, text, flags=re.MULTILINE)
    return re.sub(r"\n{3,}","\n\n",text).strip()

def parse_html_blocks(body_html):
    blocks=[]; html=re.sub(r"\s+"," ",body_html); table_data=[]
    def grab(m):
        raw=m.group(0)
        headers=[re.sub(r"<[^>]+>","",h).strip() for h in re.findall(r"<th[^>]*>(.*?)</th>",raw,re.DOTALL)]
        rows=[]
        for rr in re.findall(r"<tr[^>]*>(.*?)</tr>",raw,re.DOTALL):
            cells=[re.sub(r"<[^>]+>","",c).strip() for c in re.findall(r"<td[^>]*>(.*?)</td>",rr,re.DOTALL)]
            if cells: rows.append(cells)
        idx=len(table_data); table_data.append({"headers":headers,"rows":rows})
        return f"@@T{idx}@@"
    html=re.sub(r"<table[^>]*>.*?</table>",grab,html,flags=re.DOTALL)
    parts=re.split(r"(<[^>]+>)",html); cur=None; buf=""
    for p in parts:
        if p.startswith("<"):
            tag=p.lower().strip("<>/").split()[0] if p.strip("<>/") else ""
            if p.startswith("</"):
                c=re.sub(r"<[^>]+>","",buf).strip()
                if cur=="h1" and c: blocks.append(("h1",c))
                elif cur=="h2" and c: blocks.append(("h2",c))
                elif cur=="li" and c: blocks.append(("li",c))
                elif cur=="p" and c:
                    for m2 in re.finditer(r"@@T(\d+)@@",c):
                        blocks.append(("table",table_data[int(m2.group(1))]))
                    cl=re.sub(r"@@T\d+@@","",c).strip()
                    if cl: blocks.append(("p",cl))
                buf=""; cur=None
            else: cur=tag; buf=""
        else: buf+=p
    return blocks

def build_docx(body_html, font_name="Arial", font_size=13):
    blocks=parse_html_blocks(body_html)
    doc=Document()
    for s in doc.sections:
        s.top_margin=s.bottom_margin=Inches(1)
        s.left_margin=s.right_margin=Inches(1.2)
    doc.styles["Normal"].font.name=font_name
    doc.styles["Normal"].font.size=Pt(font_size*0.75)
    for kind,content in blocks:
        if kind=="h1": add_banner(doc,content)
        elif kind=="h2":
            p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(10)
            r=p.add_run(content); r.bold=True
            r.font.size=Pt(font_size*0.75+2); r.font.name=font_name
            r.font.color.rgb=RGBColor(108,71,255)
        elif kind=="li":
            p=doc.add_paragraph(style="List Bullet")
            r=p.add_run(content); r.font.name=font_name; r.font.size=Pt(font_size*0.75)
        elif kind=="table":
            add_docx_table(doc,content["headers"],content["rows"],font_name,font_size)
        else:
            p=doc.add_paragraph(); r=p.add_run(content)
            r.font.name=font_name; r.font.size=Pt(font_size*0.75)
    return doc

# ── Routes ────────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    ui = os.path.join(os.path.dirname(os.path.abspath(__file__)), "ui.html")
    if not os.path.exists(ui):
        return "<h2>ui.html not found</h2>", 404
    return open(ui, encoding="utf-8").read()

# API key — stored in browser session cookie (private per user, never on server)
@app.route("/api/key", methods=["GET"])
def get_key():
    return jsonify({"key": session.get("groq_key","")})

@app.route("/api/key", methods=["POST"])
def save_key():
    session["groq_key"] = request.json.get("key","").strip()
    session.permanent = True
    return jsonify({"ok": True})

# Settings — stored in session cookie
@app.route("/api/settings", methods=["GET"])
def get_settings():
    s = dict(DEFAULT_SETTINGS)
    s.update(session.get("settings", {}))
    return jsonify(s)

@app.route("/api/settings", methods=["POST"])
def save_settings():
    session["settings"] = request.json
    return jsonify({"ok": True})

# Session text — stored in session cookie (limited size, that's fine)
@app.route("/api/session", methods=["GET"])
def get_session_text():
    return jsonify({"text": session.get("last_text","")})

@app.route("/api/session", methods=["POST"])
def save_session_text():
    # Only keep last 10,000 chars to stay within cookie limits
    session["last_text"] = request.json.get("text","")[:10000]
    return jsonify({"ok": True})

@app.route("/api/recent", methods=["GET"])
def get_recent():
    return jsonify([])   # No filesystem on cloud — recent not applicable

# ── Streaming format ──────────────────────────────────────────────────────────
@app.route("/api/format", methods=["POST"])
def format_notes():
    data     = request.json or {}
    raw_text = data.get("text","").strip()
    style    = data.get("style","Clean")
    sett     = data.get("settings", DEFAULT_SETTINGS)

    # Key from request body (sent by client) or session
    api_key = data.get("apiKey","") or session.get("groq_key","")

    if not api_key:
        def no_key():
            yield "data: "+json.dumps({"error":"No API key. Enter your free Groq key in the app."})+ "\n\n"
        return Response(no_key(), mimetype="text/event-stream")

    cleaned = clean_markdown(raw_text)
    style_desc = {
        "Clean":      "clean readable prose, clear headings, tidy bullet points",
        "Minimal":    "ultra-minimal, small uppercase labels, dash-style bullets",
        "Structured": "bold coloured headers, compact structured layout",
    }.get(style,"clean readable prose")

    prompt = f"""You are a professional notes formatter. Convert ALL of the following notes into well-structured HTML.
Style: "{style}" — {style_desc}
RULES: Format ENTIRE doc. Strip all markdown symbols. Use h1/h2/p/ul/li/table only. Return raw HTML only — no backticks.

Notes:
{cleaned}"""

    def generate():
        try:
            client = Groq(api_key=api_key)
            with client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role":"user","content":prompt}],
                stream=True, max_tokens=8000,
            ) as stream:
                for chunk in stream:
                    token = chunk.choices[0].delta.content or ""
                    if token:
                        yield "data: "+json.dumps({"token":token})+"\n\n"
            yield "data: "+json.dumps({"done":True})+"\n\n"
        except Exception as e:
            msg=str(e)
            if "401" in msg or "api_key" in msg.lower(): msg="Invalid Groq key. Check and re-enter it."
            elif "429" in msg: msg="Rate limit hit. Wait a moment and try again."
            elif "connect" in msg.lower(): msg="Connection error. Check your internet."
            yield "data: "+json.dumps({"error":msg})+"\n\n"

    return Response(generate(), mimetype="text/event-stream",
                    headers={"Cache-Control":"no-cache","X-Accel-Buffering":"no"})

# ── Downloads ─────────────────────────────────────────────────────────────────
@app.route("/api/download-docx", methods=["POST"])
def download_docx():
    data = request.json or {}
    try:
        doc = build_docx(data.get("html",""),
                         data.get("font","Arial"), int(data.get("size",13)))
        buf = io.BytesIO(); doc.save(buf); buf.seek(0)
        return send_file(buf, as_attachment=True,
                         download_name="formatted_notes.docx",
                         mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        return jsonify({"error":str(e)}),500

@app.route("/api/download-html", methods=["POST"])
def download_html():
    html = request.json.get("html","")
    buf = io.BytesIO(html.encode("utf-8")); buf.seek(0)
    return send_file(buf, as_attachment=True,
                     download_name="formatted_notes.html", mimetype="text/html")

# ── Health check (Render uses this) ──────────────────────────────────────────
@app.route("/health")
def health():
    return "OK", 200

if __name__ == "__main__":
    import socket
    try:
        s=socket.socket(socket.AF_INET,socket.SOCK_DGRAM); s.connect(("8.8.8.8",80))
        local_ip=s.getsockname()[0]; s.close()
    except: local_ip="localhost"
    port=5050
    print(f"\n  Notes Formatter Pro — by Mr Wise")
    print(f"  PC:    http://localhost:{port}")
    print(f"  Phone: http://{local_ip}:{port}  (same WiFi)\n")
    app.run(host="0.0.0.0", port=port, debug=False, threaded=True)
